from pathlib import Path
from datetime import date
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\vitoriga\OneDrive\Desktop\CourseCore\coursecore")
OUT_DIR = ROOT / "docx"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "CourseCore_刷题与知识库_完整设计理念.docx"

GENERATED_CONCEPT = Path(r"C:\Users\vitoriga\.codex\generated_images\019fca74-bae0-7141-b7ce-bdc063965bd7\exec-7c68c7a0-8a42-4608-965d-d324d6ff4493.png")

# standard_business_brief token map + named CourseCore brand override
FONT = "Aptos"
FONT_CN = "Microsoft YaHei"
INK = "172033"
MUTED = "667085"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GREEN = "12B76A"
GREEN_DARK = "087443"
GREEN_PALE = "E8F7EF"
VIOLET = "6E5AE6"
VIOLET_PALE = "F0EDFF"
GOLD = "B7791F"
GOLD_PALE = "FFF6DD"
RED = "B42318"
RED_PALE = "FEECEB"
GRAY_FILL = "F2F4F7"
LINE = "D9DEE7"
WHITE = "FFFFFF"


def set_run_font(run, size=None, color=None, bold=None, italic=None, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CN)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=LINE, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_paragraph_spacing(p, before=0, after=6, line=1.10):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, align=None, before=0, after=6, line=1.10, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    set_paragraph_spacing(p, before, after, line)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_paragraph_spacing(p, 0, 4, 1.167)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_paragraph_spacing(p, 0, 4, 1.167)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph_spacing(p, {1: 16, 2: 12, 3: 8}[level], {1: 8, 2: 6, 3: 4}[level], 1.10)
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, color=GREEN_DARK, bold=True)
    elif level == 2:
        set_run_font(r, size=13, color=BLUE, bold=True)
    else:
        set_run_font(r, size=12, color=DARK_BLUE, bold=True)
    return p


def add_callout(doc, label, text, fill=GREEN_PALE, label_color=GREEN_DARK):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=label_color, size="10")
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, 2, 3, 1.10)
    r = p.add_run(label + "  ")
    set_run_font(r, size=10.5, color=label_color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    add_para(doc, "", size=2, after=2)
    return table


def add_table(doc, headers, rows, widths, header_fill=GRAY_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, 2, 2, 1.10)
        r = p.add_run(h)
        set_run_font(r, size=10, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_border(cell)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, 2, 2, 1.10)
            r = p.add_run(str(value))
            set_run_font(r, size=10, color=INK)
    add_para(doc, "", size=2, after=2)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, 2, 8, 1.0)
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED, italic=True)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=9, color=MUTED)


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Title", 25, INK, 0, 4),
        ("Subtitle", 13, MUTED, 0, 16),
        ("Heading 1", 16, GREEN_DARK, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.10

    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        st.font.size = Pt(11)
        st.font.color.rgb = RGBColor.from_string(INK)
        st.paragraph_format.left_indent = Inches(0.5)
        st.paragraph_format.first_line_indent = Inches(-0.25)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.167

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "CourseCore  |  刷题与知识库设计理念"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(hp, 0, 0, 1.0)
    for run in hp.runs:
        set_run_font(run, size=8.5, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_title_page(doc):
    add_para(doc, "PRODUCT DESIGN PRINCIPLES", size=10, color=GREEN_DARK, bold=True, after=12)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 8, 1.0)
    r = p.add_run("CourseCore")
    set_run_font(r, size=30, color=INK, bold=True)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 8, 1.0)
    r = p.add_run("刷题与知识库完整设计理念")
    set_run_font(r, size=23, color=GREEN_DARK, bold=True)
    add_para(doc, "从题目工具到持续提升的学习系统", size=14, color=MUTED, after=22)
    add_callout(doc, "核心命题", "不把 CourseCore 做成题目展示工具，而是做成一个能帮助用户持续提高的学习闭环。", fill=GREEN_PALE, label_color=GREEN_DARK)
    add_para(doc, "", size=3, after=8)
    meta = add_table(doc, ["文档属性", "内容"], [
        ("版本", "V1.0"),
        ("日期", "2026 年 8 月 4 日"),
        ("适用范围", "刷题首页、试卷流程、成绩总结、错题库、知识库、社区内容"),
        ("设计基线", "深色界面、绿色主行动色、12 栏栅格、任务驱动的信息架构"),
    ], [2300, 7060], header_fill=VIOLET_PALE)
    add_para(doc, "", size=3, after=4)
    add_para(doc, "文档使用方式", size=11, color=GREEN_DARK, bold=True, after=4)
    add_para(doc, "本文件用于产品、设计和开发共同评审。它解释“为什么这样设计”，同时把理念落到页面结构、功能取舍和验收标准上。", size=11, color=INK, after=6)
    doc.add_page_break()


def add_body(doc):
    add_heading(doc, "一、设计总纲：围绕持续提升，而不是围绕功能堆叠", 1)
    add_para(doc, "CourseCore 已经具备刷题、试卷、错题、收藏、知识库和社区等功能。下一阶段的重点不是继续增加入口，而是把这些入口组织成用户能理解、能完成、能复习的学习路径。", after=6)
    add_callout(doc, "设计判断", "用户并不缺少“更多内容”，用户缺少的是：针对自己的问题，下一步最值得做什么。", fill=GOLD_PALE, label_color=GOLD)
    add_para(doc, "因此，所有页面都遵循“状态 → 诊断 → 行动”的基本结构：先告诉用户当前状态，再解释问题来源，最后给出明确行动。", after=6)

    add_heading(doc, "二、用户模型：同一个人，在不同阶段有不同任务", 1)
    add_table(doc, ["学习阶段", "用户问题", "界面应提供"], [
        ("开始前", "我今天应该做什么？", "继续上次练习、今日目标、预计时长"),
        ("刷题中", "这道题怎么判断？我能否快速继续？", "清晰题干、即时反馈、标记、跳过、解析"),
        ("完成后", "我为什么错？哪里最需要补？", "成绩诊断、错误原因、薄弱知识点"),
        ("复习时", "怎样避免再次出错？", "知识点、相似题、复习间隔、掌握度"),
    ], [1700, 3650, 4010])
    add_para(doc, "这也是为什么首页主入口应从“添加试卷”改为“继续上次练习”：它更贴近用户当前任务，也更容易形成连续学习。", after=6)

    add_heading(doc, "三、五条核心设计理念", 1)
    add_heading(doc, "1. 以用户任务为中心，而不是以功能为中心", 2)
    add_para(doc, "产品导航可以按模块组织，但首屏内容应按任务组织。用户需要的是“继续学习、今日复习、解决薄弱点”，而不是先理解系统拥有多少模块。", after=6)
    add_bullet(doc, "刷题首页优先展示继续练习和今日任务。")
    add_bullet(doc, "排行榜、添加试卷等功能降为次级入口。")
    add_bullet(doc, "每个页面最多保留一个主要行动按钮。")

    add_heading(doc, "2. 让数据最终转化为行动", 2)
    add_para(doc, "正确率本身不是结论。70% 只说明结果，只有在它被解释为“极限与连续掌握度 42%，建议复习 6 道错题”时，数据才真正帮助用户。", after=6)
    add_bullet(doc, "总结页必须同时展示表现、问题和下一步。")
    add_bullet(doc, "知识点分析优先使用可比较的进度条、排序和趋势。")
    add_bullet(doc, "每一个重要数据旁边尽量提供对应动作。")

    add_heading(doc, "3. 建立刷题、错题和知识库的闭环", 2)
    add_para(doc, "错题不是终点，知识库也不是内容仓库。两者必须通过知识点和相似题连接起来，形成重复验证的学习循环。", after=6)
    add_table(doc, ["触发场景", "系统动作", "用户收益"], [
        ("答题错误", "记录错误原因和关联知识点", "知道问题属于哪里"),
        ("打开错题", "提供知识点、解析和相似题", "知道为什么错、怎么练"),
        ("完成复习", "更新掌握度和下次复习时间", "知道是否真正掌握"),
    ], [2200, 3900, 3260])

    add_heading(doc, "4. 用可比较的信息替代装饰性信息", 2)
    add_para(doc, "雷达图适合表达形状，但不适合支撑复习决策。用户通常需要比较“哪个知识点最低、错了多少次、多久没复习”，因此应以进度条、排序、趋势和动作按钮为主，雷达图作为辅助。", after=6)
    add_bullet(doc, "掌握度：进度条 + 百分比。")
    add_bullet(doc, "问题严重程度：错题数 + 最近错误时间。")
    add_bullet(doc, "学习变化：较上次提升或下降。")
    add_bullet(doc, "下一步：直接提供复习、巩固或练习相似题。")

    add_heading(doc, "5. 用设计系统减少随意感", 2)
    add_para(doc, "当前视觉风格已经有深色背景和绿色操作色的识别基础。改版不需要推翻，而是通过栅格、间距、组件和状态规范，让所有页面看起来属于同一个产品。", after=6)
    add_table(doc, ["设计对象", "建议规则"], [
        ("栅格", "桌面端采用 12 栏布局，内容最大宽度 1200–1360px"),
        ("间距", "以 8px 为基础单位，页面常用 16 / 24 / 32px"),
        ("主色", "绿色只用于主要行动、进度和成功状态"),
        ("状态", "颜色配合文字或图标，不单独依赖颜色"),
        ("卡片", "统一标题、数据、操作三段式结构，减少孤立装饰卡片"),
        ("公式", "所有题干、选项和解析使用统一数学排版"),
    ], [2200, 7160], header_fill=GREEN_PALE)

    add_heading(doc, "四、信息架构：把模块组织成学习路径", 1)
    add_para(doc, "建议保留刷题、知识库和社区的模块边界，但在任务层面打通它们。用户可以从任意入口进入同一条学习闭环。", after=6)
    add_table(doc, ["模块", "核心任务", "关键入口"], [
        ("今日学习", "知道今天该做什么", "继续上次练习、今日目标、待复习"),
        ("刷题", "完成题目并获得反馈", "按试卷、按题型、按知识点"),
        ("总结", "理解成绩和问题", "成绩、错误原因、薄弱点、下一步"),
        ("知识库", "补足概念并建立记忆", "知识点、错题、收藏、复习计划"),
        ("社区", "获取他人的方法和经验", "文章、讨论、收藏到知识库"),
    ], [1700, 3200, 4460])
    add_callout(doc, "闭环", "做题 → 发现错误 → 定位知识点 → 查看解析 → 练习相似题 → 再次验证掌握度。", fill=VIOLET_PALE, label_color=VIOLET)

    add_heading(doc, "五、页面设计理念落地", 1)
    add_heading(doc, "5.1 刷题首页：先解决“现在做什么”", 2)
    add_para(doc, "刷题首页不应该首先展示排行榜和添加试卷，而应该先展示个人学习任务。主卡片应包括试卷名称、当前进度、预计剩余时间、薄弱章节和继续刷题按钮。", after=6)
    add_bullet(doc, "首屏主卡片：继续上次练习。")
    add_bullet(doc, "首屏次卡片：今日待复习和即将遗忘的知识点。")
    add_bullet(doc, "第二层：薄弱知识点和最近成绩趋势。")
    add_bullet(doc, "第三层：我的试卷、按题型练习和自定义组卷。")

    add_heading(doc, "5.2 刷题总结页：从“报告成绩”到“指导下一次”", 2)
    add_para(doc, "总结页的价值不在于占满空间，而在于让用户快速理解差距和行动。建议采用四层结构：本次表现、问题诊断、错题清单、推荐下一步。", after=6)
    add_table(doc, ["层级", "展示内容", "设计目的"], [
        ("表现", "得分、答题情况、用时、较上次变化", "建立结果认知"),
        ("诊断", "知识点、错误原因、题型、用时异常", "解释结果来源"),
        ("错题", "题号、章节、错误原因、复习状态", "定位具体问题"),
        ("行动", "立即复习、查看知识点、练相似题", "推动下一步学习"),
    ], [1500, 4350, 3510])

    add_heading(doc, "5.3 知识库：从“内容存放”到“复习控制台”", 2)
    add_para(doc, "知识库首页要让用户快速知道哪些内容等待处理、哪些知识点正在遗忘、哪一部分掌握度最低。知识点列表使用可排序进度表，比单独使用雷达图更适合复习决策。", after=6)
    add_bullet(doc, "今日待复习题数。")
    add_bullet(doc, "即将遗忘的知识点。")
    add_bullet(doc, "当前最薄弱的三个章节。")
    add_bullet(doc, "最近掌握度变化。")
    add_bullet(doc, "统一搜索知识点、错题和文章。")

    add_heading(doc, "5.4 错题复习：回答“为什么错”", 2)
    add_para(doc, "错题复习页需要明确区分用户答案和正确答案，并把错误原因、核心解析、关联知识点和相似题放在同一页，避免用户在多个页面之间跳转。", after=6)
    add_table(doc, ["信息", "推荐表达"], [
        ("答案", "你的答案：B / 正确答案：A"),
        ("错误原因", "概念不清、计算错误、审题错误、方法不熟、时间不足"),
        ("解析", "先给一句结论，再展开步骤和公式"),
        ("关联内容", "关联知识点、相似题、复习计划"),
        ("完成状态", "标记为已掌握，并记录下一次复习时间"),
    ], [2200, 7160], header_fill=RED_PALE)

    add_heading(doc, "六、视觉语言与情绪设计", 1)
    add_para(doc, "深色主题适合长时间学习场景，但必须避免低对比度和过度沉重。绿色的作用不是装饰，而是告诉用户“这里可以继续、这里已经掌握、这里是下一步”。", after=6)
    add_table(doc, ["视觉元素", "心理作用", "使用边界"], [
        ("深色背景", "降低视觉噪音，突出内容", "保证次要文字对比度"),
        ("绿色主色", "传达进展、成功和可继续", "不要把所有按钮都染成绿色"),
        ("紫色辅助色", "区分筛选、题型和辅助状态", "不与成功状态混用"),
        ("红色错误色", "提醒用户存在需要处理的问题", "同时配合文字，避免只依赖颜色"),
        ("适度圆角", "降低工具感，提升亲和力", "统一圆角等级，减少每块都像卡片"),
    ], [2300, 3500, 3560])
    add_callout(doc, "情绪目标", "用户完成一次练习后，即使成绩不理想，也应该觉得问题是清楚的、可拆解的、下一步可执行的。", fill=GREEN_PALE, label_color=GREEN_DARK)

    add_heading(doc, "七、功能取舍：丰富，但不让用户迷路", 1)
    add_para(doc, "功能丰富不等于入口越多越好。新功能应该被放在用户需要它的时刻，而不是全部堆在首页。", after=6)
    add_table(doc, ["优先级", "能力", "原因"], [
        ("P0", "公式渲染、首页主任务、总结行动区、错题答案结构", "直接影响学习效率和可信度"),
        ("P1", "错题—知识点双向关联、相似题、遗忘风险、复习计划", "建立学习闭环"),
        ("P2", "AI 解析、专项练习、社区收藏、成就体系", "增强差异化和长期留存"),
    ], [1200, 4700, 3460], header_fill=GOLD_PALE)
    add_para(doc, "排行榜和游戏化功能可以保留，但应作为辅助激励，不应取代用户自己的目标和进步。", after=6)

    add_heading(doc, "八、用户中心原则：少打断，多反馈", 1)
    add_heading(doc, "学习开始前", 2)
    add_bullet(doc, "提供继续上次学习、今日目标和预计时长。")
    add_bullet(doc, "默认推荐当前最值得做的内容，同时允许用户自定义。")
    add_heading(doc, "学习进行中", 2)
    add_bullet(doc, "自动保存题目进度，允许跳过和返回。")
    add_bullet(doc, "解析、知识点和标记操作尽量不离开当前页面。")
    add_bullet(doc, "使用文字、图标和状态标签共同表达结果。")
    add_heading(doc, "学习完成后", 2)
    add_bullet(doc, "展示成绩变化，而不是只展示绝对分数。")
    add_bullet(doc, "告诉用户最薄弱的三个知识点。")
    add_bullet(doc, "给出明确、可执行的下一步学习动作。")

    add_heading(doc, "九、验收标准", 1)
    add_para(doc, "设计理念最终需要落实为可检查的产品行为，建议使用以下标准验收：", after=6)
    checks = [
        "用户进入刷题首页后，3 秒内能找到继续学习入口。",
        "用户完成试卷后，首屏可以看到成绩、问题诊断和下一步动作。",
        "任意一道错题都能跳转到对应知识点。",
        "任意一个知识点都能看到关联错题和相似题。",
        "公式、题干、选项和解析在所有页面均能正确排版。",
        "状态不单独依赖颜色表达，同时有文字或图标说明。",
        "页面在 1280px 以上桌面宽度下保持稳定栅格和合理密度。",
        "新增功能不会抢占用户主任务的视觉优先级。",
    ]
    for item in checks:
        add_bullet(doc, item)

    add_heading(doc, "十、实施建议", 1)
    add_para(doc, "推荐按“先修体验底座，再打通学习闭环，最后增加差异化能力”的顺序实施。", after=6)
    for item in [
        "第一阶段：修复数学公式、状态逻辑、选项布局和大面积空白。",
        "第二阶段：重做刷题首页和总结页的信息层级。",
        "第三阶段：打通错题与知识点的关联，增加相似题和复习计划。",
        "第四阶段：完善收藏、社区、AI 解析和激励体系。",
    ]:
        add_number(doc, item)
    add_callout(doc, "最终目标", "让用户在每个页面都清楚地知道：我现在在哪里、我的问题是什么、下一步该做什么。", fill=VIOLET_PALE, label_color=VIOLET)

    if GENERATED_CONCEPT.exists():
        doc.add_page_break()
        add_heading(doc, "附录：视觉方向参考", 1)
        add_para(doc, "以下视觉参考保留 CourseCore 当前的深色背景与绿色行动色，同时强调首屏任务、知识点进度和学习趋势。它用于说明视觉气质，不替代页面结构和交互规范。", after=8)
        doc.add_picture(str(GENERATED_CONCEPT), width=Inches(6.3))
        add_caption(doc, "图 1  CourseCore 刷题首页视觉方向参考")


def main():
    doc = Document()
    configure_doc(doc)
    add_title_page(doc)
    add_body(doc)
    doc.core_properties.title = "CourseCore 刷题与知识库完整设计理念"
    doc.core_properties.subject = "CourseCore 产品设计理念与学习闭环"
    doc.core_properties.author = "CourseCore Design"
    doc.core_properties.comments = "Generated design rationale document"
    doc.save(str(OUT))
    print(str(OUT))


if __name__ == "__main__":
    main()
